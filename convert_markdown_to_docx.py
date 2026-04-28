#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import re


def parse_markdown_to_docx(input_path, output_path):
    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        text = line.strip()

        if not text:
            doc.add_paragraph()
            i += 1
            continue

        if text.startswith('# '):
            heading = text[2:].strip()
            p = doc.add_heading(heading, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if text.startswith('## '):
            heading = text[3:].strip()
            doc.add_heading(heading, level=2)
            i += 1
            continue

        if text.startswith('### '):
            heading = text[4:].strip()
            doc.add_heading(heading, level=3)
            i += 1
            continue

        if text.startswith('#### '):
            heading = text[5:].strip()
            doc.add_heading(heading, level=4)
            i += 1
            continue

        if text.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            doc.add_paragraph(''.join(code_lines))
            i += 1
            continue

        if re.match(r'^\d+\.\s+', text):
            content = re.sub(r'^\d+\.\s+', '', text)
            doc.add_paragraph(content, style='List Number')
            i += 1
            continue

        if text.startswith('- '):
            content = text[2:].strip()
            doc.add_paragraph(content, style='List Bullet')
            i += 1
            continue

        if text.startswith('  - '):
            content = text[4:].strip()
            doc.add_paragraph(content, style='List Bullet 2')
            i += 1
            continue

        if '|' in text and re.search(r'\|', text):
            doc.add_paragraph(text)
            i += 1
            continue

        clean = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        clean = re.sub(r'\*(.*?)\*', r'\1', clean)
        clean = re.sub(r'`(.*?)`', r'\1', clean)
        clean = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', clean)
        doc.add_paragraph(clean)
        i += 1

    doc.save(output_path)
    print(f"✅ Created Word document: {output_path}")


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print('Usage: python convert_markdown_to_docx.py input.md output.docx')
        sys.exit(1)
    parse_markdown_to_docx(sys.argv[1], sys.argv[2])
