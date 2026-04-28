#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convert Markdown Report to Word Document
Vietnamese encoding support
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_word_document():
    """Create Word document from markdown report"""
    
    # Read markdown file
    with open('BAOCAO_DOANHOM.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Parse and process content
    lines = content.split('\n')
    
    for line in lines:
        # Skip empty lines
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # Headers (# = Heading 1, ## = Heading 2, etc.)
        if line.startswith('# '):
            heading = line[2:].strip()
            p = doc.add_heading(heading, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        elif line.startswith('## '):
            heading = line[3:].strip()
            p = doc.add_heading(heading, level=2)
        
        elif line.startswith('### '):
            heading = line[4:].strip()
            p = doc.add_heading(heading, level=3)
        
        elif line.startswith('#### '):
            heading = line[5:].strip()
            p = doc.add_heading(heading, level=4)
        
        # Unordered lists
        elif line.startswith('- '):
            text = line[2:].strip()
            doc.add_paragraph(text, style='List Bullet')
        
        elif line.startswith('  - '):
            text = line[4:].strip()
            doc.add_paragraph(text, style='List Bullet 2')
        
        # Ordered lists
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line).strip()
            doc.add_paragraph(text, style='List Number')
        
        # Code blocks (detect by indent or backticks)
        elif line.startswith('```'):
            # Find code block end
            idx = lines.index(line)
            code_lines = []
            i = idx + 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.style = 'Normal'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 0, 128)
        
        # Tables (detect by pipes |)
        elif '|' in line and '-' in line:
            # Skip for now, will handle separately
            doc.add_paragraph(line)
        
        # Regular paragraphs
        else:
            # Remove markdown formatting
            text = line.strip()
            # Remove bold markers
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            # Remove italic markers
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            # Remove inline code markers
            text = re.sub(r'`(.*?)`', r'\1', text)
            # Remove links
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
            
            if text:
                doc.add_paragraph(text)
    
    # Save document
    doc.save('BAOCAO_DOANHOM.docx')
    print("✅ Word document created: BAOCAO_DOANHOM.docx")

if __name__ == '__main__':
    create_word_document()
